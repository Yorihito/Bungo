"""
Tests for bungo.converter (to_text, to_docx, convert).
"""

from __future__ import annotations

import pytest

from bungo.converter import convert, to_docx, to_text
from bungo.parser import BungoDocument


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_doc(paragraphs: list[str]) -> BungoDocument:
    return BungoDocument(paragraphs=paragraphs)


# ---------------------------------------------------------------------------
# to_text
# ---------------------------------------------------------------------------


class TestToText:
    def test_creates_file(self, tmp_path):
        doc = _make_doc(["Hello", "World"])
        out = tmp_path / "output.txt"
        to_text(doc, str(out))
        assert out.exists()

    def test_content_preserved(self, tmp_path):
        doc = _make_doc(["First line", "Second line"])
        out = tmp_path / "output.txt"
        to_text(doc, str(out))
        content = out.read_text(encoding="utf-8")
        assert "First line" in content
        assert "Second line" in content

    def test_japanese_content(self, tmp_path):
        doc = _make_doc(["日本語テキスト", "文豪miniの文書"])
        out = tmp_path / "output.txt"
        to_text(doc, str(out))
        content = out.read_text(encoding="utf-8")
        assert "日本語テキスト" in content
        assert "文豪miniの文書" in content

    def test_empty_document(self, tmp_path):
        doc = _make_doc([])
        out = tmp_path / "output.txt"
        to_text(doc, str(out))
        assert out.exists()

    def test_encoding_is_utf8(self, tmp_path):
        doc = _make_doc(["テスト"])
        out = tmp_path / "output.txt"
        to_text(doc, str(out))
        # Should decode without error using UTF-8
        content = out.read_bytes().decode("utf-8")
        assert "テスト" in content


# ---------------------------------------------------------------------------
# to_docx
# ---------------------------------------------------------------------------


class TestToDocx:
    def test_creates_file(self, tmp_path):
        doc = _make_doc(["Hello", "World"])
        out = tmp_path / "output.docx"
        to_docx(doc, str(out))
        assert out.exists()

    def test_docx_content_preserved(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document as DocxDocument  # type: ignore[import-untyped]

        doc = _make_doc(["First paragraph", "Second paragraph"])
        out = tmp_path / "output.docx"
        to_docx(doc, str(out))

        word_doc = DocxDocument(str(out))
        texts = [p.text for p in word_doc.paragraphs]
        joined = "\n".join(texts)
        assert "First paragraph" in joined
        assert "Second paragraph" in joined

    def test_docx_japanese_content(self, tmp_path):
        pytest.importorskip("docx")
        from docx import Document as DocxDocument  # type: ignore[import-untyped]

        doc = _make_doc(["日本語テキスト"])
        out = tmp_path / "output.docx"
        to_docx(doc, str(out))

        word_doc = DocxDocument(str(out))
        texts = [p.text for p in word_doc.paragraphs]
        assert any("日本語テキスト" in t for t in texts)

    def test_creates_parent_dirs(self, tmp_path):
        doc = _make_doc(["Hello"])
        out = tmp_path / "subdir" / "nested" / "output.docx"
        to_docx(doc, str(out))
        assert out.exists()


# ---------------------------------------------------------------------------
# convert dispatch
# ---------------------------------------------------------------------------


class TestConvert:
    def test_dispatch_txt(self, tmp_path):
        doc = _make_doc(["Hello"])
        out = tmp_path / "output.txt"
        convert(doc, str(out), "txt")
        assert out.exists()
        assert "Hello" in out.read_text(encoding="utf-8")

    def test_dispatch_docx(self, tmp_path):
        doc = _make_doc(["Hello"])
        out = tmp_path / "output.docx"
        convert(doc, str(out), "docx")
        assert out.exists()

    def test_dispatch_with_dot_prefix(self, tmp_path):
        """Format strings like '.txt' and '.docx' should also work."""
        doc = _make_doc(["Hello"])
        out = tmp_path / "output.txt"
        convert(doc, str(out), ".txt")
        assert out.exists()

    def test_unknown_format_raises(self, tmp_path):
        doc = _make_doc(["Hello"])
        out = tmp_path / "output.pdf"
        with pytest.raises(ValueError, match="Unsupported output format"):
            convert(doc, str(out), "pdf")
