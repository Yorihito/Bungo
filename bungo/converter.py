"""
Output converters for parsed 文豪mini documents.

Supported output formats:

* **txt** – Plain-text UTF-8 file.
* **docx** – Microsoft Word Open XML document (via *python-docx*).
"""

from __future__ import annotations

import io
from pathlib import Path
from typing import Union

from .parser import BungoDocument

PathLike = Union[str, Path]


# ---------------------------------------------------------------------------
# Plain text
# ---------------------------------------------------------------------------


def to_text(doc: BungoDocument, output_path: PathLike) -> None:
    """Write *doc* as a plain-text file encoded in UTF-8.

    Ruby (annotations) are filtered out for better readability.

    Args:
        doc: Parsed :class:`~bungo.parser.BungoDocument`.
        output_path: Destination file path (created or overwritten).
    """
    lines = []
    for para in doc.paragraphs:
        # Filter out ruby segments
        text = "".join(s.text for s in para if not s.is_ruby)
        lines.append(text)
    
    output = "\n".join(lines)
    Path(output_path).write_text(output, encoding="utf-8")


# ---------------------------------------------------------------------------
# Microsoft Word (.docx)
# ---------------------------------------------------------------------------


def to_docx(doc: BungoDocument, output_path: PathLike) -> None:
    """Write *doc* as a Microsoft Word (*.docx*) file.

    Includes all segments (including Ruby) as they appeared in the stream.
    Uses a fixed-width font and smaller text for Ruby paragraphs to
    simulate the original word processor layout.

    Args:
        doc: Parsed :class:`~bungo.parser.BungoDocument`.
        output_path: Destination file path (created or overwritten).

    Raises:
        ImportError: If *python-docx* is not installed.
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
        from docx.oxml.ns import qn
        from docx.shared import Pt
    except ImportError as exc:  # pragma: no cover
        raise ImportError(
            "python-docx is required for Word output.  "
            "Install it with: pip install python-docx"
        ) from exc

    word_doc = Document()

    # Set default font to MS Gothic (standard monospace Japanese font)
    style = word_doc.styles["Normal"]
    font = style.font
    font.name = "MS Gothic"
    font.size = Pt(10.5)
    # Set East Asia font (required for CJK fonts in docx)
    r = font._element.get_or_add_rPr()
    rFonts = r.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), "MS Gothic")

    for para in doc.paragraphs:
        # Detect if it's a Ruby-only paragraph (all non-space segments are ruby)
        is_ruby_para = all(s.is_ruby or s.text.isspace() for s in para) if para else False

        para_text = "".join(s.text for s in para)
        p = word_doc.add_paragraph()
        run = p.add_run(para_text)

        # Basic alignment/spacing settings for word-processor feel
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        if is_ruby_para:
            # Use gray color to distinguish ruby without breaking alignment.
            # Using the same size (10.5pt) ensures that spaces and characters
            # align perfectly with the line below.
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Ensure parent directory exists.
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    word_doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Convenience helper used by the CLI
# ---------------------------------------------------------------------------


def convert(doc: BungoDocument, output_path: PathLike, fmt: str) -> None:
    """Dispatch to the appropriate converter based on *fmt*.

    Args:
        doc: Parsed document.
        output_path: Destination path.
        fmt: ``"txt"`` or ``"docx"``.

    Raises:
        ValueError: If *fmt* is not a recognised format name.
    """
    fmt = fmt.lower().lstrip(".")
    if fmt == "txt":
        to_text(doc, output_path)
    elif fmt == "docx":
        to_docx(doc, output_path)
    else:
        raise ValueError(f"Unsupported output format: {fmt!r}.  Choose 'txt' or 'docx'.")
